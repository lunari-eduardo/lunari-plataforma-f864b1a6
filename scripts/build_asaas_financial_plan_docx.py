from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from datetime import date
from pathlib import Path

OUT = Path("docs/PLANO_IMPLEMENTACAO_FINANCEIRO_ASAAS.docx")

NAVY = "17324D"
BLUE = "2E74B5"
MUTED = "5B6673"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
GREEN = "E7F4EA"
AMBER = "FFF4DB"
WHITE = "FFFFFF"
BLACK = "1A1A1A"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, inches):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is not None:
        tblW.set(qn("w:w"), str(int(sum(widths) * 1440)))
        tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for gc, width in zip(grid.gridCol_lst, widths):
        gc.set(qn("w:w"), str(int(width * 1440)))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margin(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, size=10.5, color=BLACK, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_para(doc, text="", size=10.5, color=BLACK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size, color, bold, italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(text), 10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    set_run(p.add_run(text), 10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    if level == 1:
        set_run(r, 15, BLUE, True)
    else:
        set_run(r, 12, NAVY, True)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(title), 10.5, NAVY, True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    set_run(p2.add_run(body), 10.2, BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for cell, text in zip(header.cells, headers):
        shade(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), 9.5, NAVY, True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            set_run(p.add_run(str(text)), 9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def page_break(doc):
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
styles["Normal"].font.size = Pt(10.5)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(header.add_run("LUNARI STUDIO  |  PLANO FINANCEIRO ASAAS"), 8.5, MUTED, True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run(footer.add_run("Documento de planejamento - versão 1.0 - 28 de agosto de 2026"), 8, MUTED)

# Opening masthead
add_para(doc, "PLANO DE IMPLEMENTAÇÃO", 10, BLUE, True, after=3)
add_para(doc, "Financeiro Asaas: vendas, contas a receber e caixa", 23, NAVY, True, after=5)
add_para(doc, "Arquitetura, correções de consistência e roteiro de execução", 13, MUTED, after=18)

meta = doc.add_table(rows=4, cols=2)
set_table_geometry(meta, [1.35, 5.15])
for i, (label, value) in enumerate([
    ("Produto", "Lunari Studio"),
    ("Escopo", "Integração Asaas, Workflow, cobranças e Financeiro"),
    ("Status", "Plano aprovado para execução por fases; nenhuma mudança aplicada neste documento"),
    ("Próxima etapa", "Fase 0 - auditoria e inventário de dados, mediante autorização"),
]):
    shade(meta.cell(i, 0), LIGHT_GRAY)
    p = meta.cell(i, 0).paragraphs[0]
    set_run(p.add_run(label), 9.4, NAVY, True)
    p2 = meta.cell(i, 1).paragraphs[0]
    set_run(p2.add_run(value), 9.4)

add_callout(doc, "DECISÃO DE ARQUITETURA", "O Workflow continuará sendo a fonte das vendas por competência. O Asaas passará a alimentar contas a receber e caixa, sem transformar taxas, antecipações ou valores de parcelamento em receita de sessão.", GREEN)

add_heading(doc, "1. Resultado esperado", 1)
add_para(doc, "O Lunari deverá apresentar três leituras financeiras distintas, mas vinculadas pela mesma venda:")
add_table(doc, ["Leitura", "Fonte canônica", "Uso"], [
    ("Venda comercial", "clientes_sessoes", "Métricas de vendas, ticket médio e faturamento por período da sessão."),
    ("Contas a receber", "cobrancas + cobranca_parcelas", "Saldo do cliente, parcelas, vencimentos e quitação."),
    ("Caixa realizado", "Novo razão de movimentos do gateway", "Créditos, tarifas, antecipações, estornos e saldo efetivo."),
], [1.35, 2.05, 3.10])

add_heading(doc, "2. Princípios obrigatórios", 1)
for item in [
    "A venda de uma sessão não é alterada pelo calendário de recebimento.",
    "Uma parcela pode estar paga pelo cliente e ainda não estar disponível em caixa.",
    "Toda tarifa é custo, recuperação de custo ou ajuste; nunca receita da sessão.",
    "Eventos e extrato do Asaas são a autoridade para valores e datas efetivos do gateway.",
    "Configurações de pagamento valem para novas cobranças e são congeladas em cada cobrança emitida.",
    "Reprocessar um webhook ou uma sincronização nunca pode criar uma segunda movimentação.",
]:
    add_bullet(doc, item)

page_break(doc)

add_heading(doc, "3. Fluxo-alvo: da venda ao dinheiro disponível", 1)
add_table(doc, ["Momento", "Venda/Workflow", "A receber", "Caixa"], [
    ("Sessão vendida por R$ 1.500", "Venda de R$ 1.500 na competência da sessão.", "Recebível principal vinculado à sessão, quando a venda estiver comercialmente confirmada.", "Nenhuma entrada."),
    ("Cobrança enviada", "Sem mudança.", "Cobrança emitida e política de taxa congelada.", "Nenhuma entrada."),
    ("Cliente escolhe 6x", "Sem mudança.", "Seis parcelas reais, com IDs, vencimentos e principal alocado.", "Nenhuma entrada."),
    ("Cliente paga uma parcela", "Atualiza indicador de quitação pelo principal pago.", "Parcela quitada pelo cliente.", "Aguardando disponibilidade, se aplicável."),
    ("Asaas disponibiliza/antecipa", "Sem mudança.", "Parcela recebe a situação de disponibilidade/antecipação.", "Crédito e taxa reais registrados."),
    ("Estorno ou chargeback", "Preserva venda, salvo decisão comercial explícita.", "Parcela é revertida ou ajustada.", "Débito/reversão registrada."),
], [1.15, 1.7, 2.0, 1.65])

add_callout(doc, "EXEMPLO", "Uma sessão de R$ 1.500 em 6x continua sendo uma única venda de R$ 1.500. As seis parcelas formam o contas a receber. O fluxo de caixa só mostra cada crédito líquido na data em que o Asaas o disponibiliza ou antecipa.", LIGHT_BLUE)

add_heading(doc, "4. Antecipação: regra de negócio e comportamento externo", 1)
add_para(doc, "A antecipação será separada em configuração da conta e fato financeiro. Ela não será mais uma opção nos painéis de geração de cobrança.")
add_table(doc, ["Situação", "Regra do Lunari", "Registro esperado"], [
    ("Antecipação automática ativa", "Configuração global da conta Asaas, gerida em Configurações de Pagamento.", "Snapshot da política na cobrança; eventos e extrato confirmam valores reais."),
    ("Antecipação automática inativa", "Lunari não pede antecipação nem adiciona previsão de taxa ao cliente.", "Parcelas seguem o cronograma normal até outro evento real."),
    ("Fotógrafo antecipa direto no Asaas", "Nunca é bloqueado, mesmo que a configuração local esteja inativa.", "Evento externo cria/atualiza antecipação, taxa e crédito reais."),
    ("Antecipação negada/cancelada", "Nenhuma alteração na venda ou no principal.", "Atualiza o status da antecipação; nenhum crédito é reconhecido."),
], [1.65, 2.5, 2.35])
add_para(doc, "Importante: se a política de repasse estava inativa quando a cobrança foi emitida, uma antecipação posterior feita no Asaas será custo real do fotógrafo. O Lunari não cobrará o cliente retroativamente.", 10.2, MUTED, italic=True, after=5)

add_heading(doc, "5. Configurações de pagamento", 1)
add_para(doc, "Consolidar a configuração Asaas em uma única fonte e remover os controles de taxa e antecipação de todos os painéis de cobrança.")
for item in [
    "Meios de pagamento habilitados e limite de parcelas.",
    "Política padrão de processamento: absorver ou repassar ao cliente.",
    "Antecipação automática Asaas: estado real, elegibilidade e acionamento global.",
    "Política padrão de repasse de antecipação, habilitada apenas quando a antecipação automática estiver ativa.",
    "Consulta de taxas para estimativa comercial, sem substituir a taxa efetivamente cobrada pelo Asaas.",
]:
    add_bullet(doc, item)

page_break(doc)

add_heading(doc, "6. Modelo de dados proposto", 1)
add_para(doc, "O objetivo é preservar dados atuais, explicitar seus papéis e acrescentar um razão auditável de gateway.")
add_table(doc, ["Componente", "Papel após a mudança", "Observações"], [
    ("clientes_sessoes", "Venda comercial", "Mantém valor da sessão, competência e métricas de vendas."),
    ("cobrancas", "Documento-pai do recebível", "Congela principal, valor cobrado, origem, política e vínculo comercial."),
    ("cobranca_parcelas", "Agenda de parcelas", "Uma linha por parcela Asaas, com IDs, vencimento, quitação e disponibilidade."),
    ("gateway_events", "Caixa de entrada idempotente", "Armazena evento bruto, ID Asaas, processamento e falhas."),
    ("gateway_anticipations", "Ciclo de vida da antecipação", "Mantém ID, parcela/cobrança, status, tarifa, líquido e datas."),
    ("gateway_cash_movements", "Razão de caixa", "Uma linha por impacto efetivo no saldo do gateway, com valor assinado."),
], [1.55, 2.2, 2.75])

add_heading(doc, "Campos financeiros explícitos", 2)
for item in [
    "valor_principal: parcela ou total que pertence à venda de serviço.",
    "valor_cobrado_cliente: total apresentado e pago pelo cliente, incluindo eventual repasse.",
    "taxa_processamento_real e taxa_antecipacao_real: valores provenientes da conciliação Asaas.",
    "valor_liquido_creditado: efeito financeiro real; não é receita comercial.",
    "fee_policy_snapshot: configuração usada quando a cobrança foi emitida.",
    "source_event_id e provider_transaction_id: chaves para rastreabilidade e deduplicação.",
]:
    add_bullet(doc, item)

add_callout(doc, "REGRA CONTÁBIL FUNCIONAL", "O valor principal é receita comercial. Taxa absorvida é despesa. Taxa repassada é recuperação de custo, apresentada separadamente. O caixa pode receber um valor líquido diferente, sem alterar a receita da sessão.", AMBER)

add_heading(doc, "7. Correções técnicas necessárias", 1)
add_table(doc, ["Problema atual", "Correção"], [
    ("Cobrança-pai marcada como paga com uma única parcela.", "Status da cobrança será derivado da soma das parcelas, sem sobrescrita pelo webhook."),
    ("Uma transação automática representa todo o parcelamento.", "Cada parcela e cada movimento de caixa terão registros próprios."),
    ("Deduplicação por tipo de evento e pagamento.", "Deduplicação pelo ID único do evento Asaas e pela transação financeira do extrato."),
    ("Taxa inferida pelo valor líquido ou intenção do usuário.", "Taxa real vem de antecipação e extrato Asaas; política apenas classifica o repasse."),
    ("Checkout relê configuração atual depois de a cobrança ser enviada.", "Checkout recebe o snapshot da cobrança, não as preferências atuais."),
    ("Parcelamento criado com campos incompatíveis.", "Usar somente installmentCount + installmentValue ou totalValue para 2+ parcelas."),
], [2.85, 3.65])

page_break(doc)

add_heading(doc, "8. Webhooks e conciliação", 1)
add_para(doc, "O webhook será rápido, autenticado, persistente e idempotente. O processamento financeiro ocorrerá a partir do evento salvo e poderá ser reexecutado com segurança.")
add_table(doc, ["Grupo", "Eventos mínimos", "Efeito"], [
    ("Cobranças", "CREATED, UPDATED, CONFIRMED, RECEIVED, OVERDUE, DELETED, RESTORED", "Atualiza cobrança, parcela e quitação do cliente."),
    ("Reversões", "REFUNDED, PARTIALLY_REFUNDED, CHARGEBACK e reversões", "Cria movimentos negativos e reabre/ajusta o recebível."),
    ("Antecipações", "RECEIVABLE_ANTICIPATION_SCHEDULED, PENDING, CREDITED, DEBITED, CANCELLED, DENIED, OVERDUE", "Atualiza antecipação e disponibilidade real de caixa."),
    ("Extrato", "GET /v3/financialTransactions", "Concilia créditos, taxas, antecipações e reversões que impactaram saldo."),
], [1.15, 3.4, 1.95])

add_heading(doc, "Regras de robustez", 2)
for item in [
    "Validar o header asaas-access-token com segredo exclusivo por integração.",
    "Persistir o payload primeiro e responder HTTP 2xx sem aguardar cálculos pesados.",
    "Permitir eventos repetidos e fora de ordem sem duplicar movimento ou regredir estado.",
    "Conciliar periodicamente cobranças, antecipações e extrato com paginação.",
    "Classificar movimentações não vinculadas ao Lunari como não conciliadas; nunca anexá-las automaticamente a uma sessão.",
]:
    add_bullet(doc, item)

add_heading(doc, "9. Fases de implementação", 1)
add_table(doc, ["Fase", "Objetivo", "Saída / portão de aprovação"], [
    ("0", "Auditoria e inventário", "Relatório de dados atuais, inconsistências e volume de migração. Nenhuma escrita de negócio."),
    ("1", "Schema compatível", "Novas tabelas/colunas e restrições, sem alterar telas ou lançamentos existentes."),
    ("2", "Emissão e ingestão", "Snapshot de política, parcelas, webhooks e antecipações com flag interna."),
    ("3", "Backfill e reconciliação", "Histórico conciliado e divergências classificadas, sem apagar legado."),
    ("4", "Corte financeiro", "Novo razão alimenta fluxo de caixa; prevenção de duplicidade ativa."),
    ("5", "Limpeza controlada", "Remoção de cálculos e controles legados após período de observação."),
], [0.55, 2.15, 3.80])

page_break(doc)

add_heading(doc, "10. Fase 0 - auditoria e inventário", 1)
add_callout(doc, "ESCOPO DA FASE 0", "A Fase 0 é somente leitura. Ela não modifica cobrança, sessão, integração, webhook, dados financeiros ou configuração do fotógrafo.", GREEN)
add_para(doc, "Objetivo: medir o estado real do banco e da integração antes de definir a migração e o corte de produção.")

add_heading(doc, "Atividades", 2)
for item in [
    "Inventariar cobranças Asaas por status, parcela, finalidade, valor, IDs de provider e período.",
    "Comparar cobrancas, cobranca_parcelas, clientes_transacoes e extrato_unificado para localizar duplicidades e lacunas.",
    "Identificar cobranças parceladas cujo status-pai diverge das parcelas.",
    "Medir transações automáticas de valor total em vendas parceladas.",
    "Identificar taxas incluídas em receitas, taxa sem origem e valores líquidos incompatíveis.",
    "Mapear todas as telas e funções que geram cobrança, incluindo checkout público, Workflow, Agenda, Financeiro e galerias.",
    "Validar integrações Asaas ativas, ambiente, webhook cadastrado, eventos assinados e cobertura de antecipação.",
    "Consultar, em amostra controlada, pagamentos, parcelas, antecipações e extrato Asaas para confirmar a estratégia de conciliação.",
]:
    add_number(doc, item)

add_heading(doc, "Entregáveis", 2)
add_table(doc, ["Entregável", "Conteúdo"], [
    ("Inventário de integridade", "Contagens, categorias de inconsistência, impacto financeiro e amostras técnicas."),
    ("Mapa de origem", "Qual tabela, função ou tela cria cada tipo de cobrança e lançamento."),
    ("Plano de migração confirmado", "Regras de backfill, exceções, ordem de implantação e estratégia de rollback."),
    ("Matriz de risco", "Dados que exigem revisão humana antes de qualquer alteração."),
], [1.85, 4.65])

add_heading(doc, "Critérios para encerrar a Fase 0", 2)
for item in [
    "Nenhuma descoberta depende de suposição não verificada sobre dados de produção.",
    "Todas as rotas de criação de cobrança estão inventariadas.",
    "Há quantificação das duplicidades e inconsistências que a migração deverá tratar.",
    "A estratégia para antecipação externa foi validada contra eventos e dados reais do Asaas.",
    "O próximo passo possui escopo, risco e plano de reversão explícitos para aprovação.",
]:
    add_bullet(doc, item)

add_heading(doc, "11. Testes e critérios de aceite", 1)
add_table(doc, ["Cenário", "Resultado obrigatório"], [
    ("Venda de R$ 1.500 em 6x", "Uma venda no Workflow; seis parcelas no receber; créditos por data real no caixa."),
    ("Antecipação feita no Asaas com configuração local inativa", "Antecipação externa e taxa real registradas, sem mudar venda ou cobrar cliente retroativamente."),
    ("Taxa repassada", "Não aumenta receita comercial; aparece como recuperação de custo separada."),
    ("Estorno parcial", "Movimento negativo único, parcela ajustada e nenhuma duplicidade."),
    ("Webhook repetido", "Nenhuma segunda parcela, antecipação ou movimentação de caixa."),
    ("Mudança futura de configuração", "Cobranças já emitidas permanecem com o snapshot original."),
], [2.55, 3.95])

add_heading(doc, "12. Referências Asaas consultadas", 1)
for text in [
    "Eventos para cobranças - https://docs.asaas.com/docs/webhook-para-cobrancas",
    "Eventos para antecipações - https://docs.asaas.com/docs/webhook-para-antecipacoes",
    "Criar uma cobrança parcelada - https://docs.asaas.com/docs/criar-uma-cobranca-parcelada",
    "Listar antecipações - https://docs.asaas.com/reference/listar-antecipacoes",
    "Recuperar extrato financeiro - https://docs.asaas.com/reference/recuperar-extrato",
    "Recuperar status da antecipação automática - https://docs.asaas.com/reference/recuperar-status-da-antecipacao-automatica",
]:
    add_bullet(doc, text)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Plano de Implementação Financeiro Asaas - Lunari Studio"
doc.core_properties.subject = "Separação entre vendas, contas a receber e fluxo de caixa"
doc.core_properties.author = "Lunari Studio"
doc.save(OUT)
print(OUT)
